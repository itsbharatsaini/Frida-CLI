from docx import Document
from mdutils.mdutils import MdUtils
from typing import List, Tuple
from fridacli.logger import Logger

logger = Logger()

def create_file(path: str, lines: List[Tuple[str, str | List]]) -> None:
    """
    Save the given lines in either .docx or .md format.

    Args:
        path (str): The path to save the document.
        lines (List[Tuple[str, str | List]]): A list of tuples containing the format and text of each line.

    Raises:
        Exception: If any error occurs during the saving process.
    """
    try:
        if "docx" in path:
            logger.info(__name__, f"(create_file) Saving recommendations (docx) in: {path}")
            doc = Document()

            for format, text in lines:
                if format == "title":
                    doc.add_heading(text)
                elif format == "subheader":
                    doc.add_heading(text, level=2)
                elif format == "bold":
                    p = doc.add_paragraph("")
                    p.add_run(text).bold = True
                elif format == "text":
                    doc.add_paragraph(text)
                elif format == "bullet_list":
                    for line in text:
                        doc.add_paragraph(line, style="List Bullet")
                elif format == "num_list":
                    for line in text:
                        doc.add_paragraph(line, style="List Number")

            doc.save(path)
        else:
            logger.info(__name__, f"(create_file) Saving recommendations (md): {path}")
            mdFile = None
            bullets = []
            for format, text in lines:
                if format == "title":
                    mdFile = MdUtils(file_name=path)
                    mdFile.new_header(level=1, title=text, add_table_of_contents="n")
                elif format == "subheader":
                    mdFile.new_header(level=2, title=text, add_table_of_contents="n")
                elif format == "bold":
                    mdFile.new_line(text, bold_italics_code="b")
                elif format == "text":
                    mdFile.new_line(text)
                elif format == "bullet_list":
                    mdFile.new_list(text)
                elif format == "num_list":
                    mdFile.new_list(text, "1")

            mdFile.create_md_file()
        logger.info(__name__, f"(create_file) Documentation saved succesfully.")
    except Exception as e:
        logger.error(__name__, f"(create_file) {e}")